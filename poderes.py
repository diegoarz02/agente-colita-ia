"""
Poderes propios de Colita.

Herramientas que no existen en ningun MCP y que necesita para su trabajo:
control del sistema, guardar en el vault, y producir entregables (Excel, HTML).

Se exponen como un servidor MCP en proceso, asi que pasan por `autoridad()`
igual que todo lo demas.
"""

from __future__ import annotations

import datetime as dt
import os
import re
import subprocess
from pathlib import Path
from typing import Any

from claude_agent_sdk import create_sdk_mcp_server, tool

VAULT = Path(r"C:\Users\diego\Documents\Obsidian Vault\Claude Diego")
SALIDAS = Path(r"C:\Users\diego\colita\salidas")
SALIDAS.mkdir(exist_ok=True)

# Sin esto, cada subproceso abre una consola negra encima de lo que estes
# haciendo. Es la causa de las "pantallas negras" que aparecen y a veces se
# quedan. Todo subprocess de este archivo lo lleva.
SIN_VENTANA = {}
if os.name == "nt":
    _si = subprocess.STARTUPINFO()
    _si.dwFlags |= subprocess.STARTF_USESHOWWINDOW
    _si.wShowWindow = 0  # SW_HIDE
    SIN_VENTANA = {
        "creationflags": subprocess.CREATE_NO_WINDOW,
        "startupinfo": _si,
    }


def _carpetas_reales() -> dict[str, Path]:
    """Dónde están de verdad Documentos y Escritorio en ESTA máquina.

    No se pueden dar por hechas: con OneDrive activado, Windows las redirige.
    En la de Diego, `Desktop` apunta a `OneDrive\\Documentos\\Datos adjuntos\\
    Desktop` y `Documents` a `OneDrive\\Documentos`, así que `~/Desktop` no
    existe. Buscarlo ahí es la razón de que Colita dijera "no encuentro nada".

    La verdad está en el registro, en User Shell Folders.
    """
    casa = Path.home()
    sitios: dict[str, Path] = {}

    if os.name == "nt":
        try:
            import winreg

            clave = r"SOFTWARE\Microsoft\Windows\CurrentVersion\Explorer\User Shell Folders"
            with winreg.OpenKey(winreg.HKEY_CURRENT_USER, clave) as k:
                for nombre, etiqueta in (("Personal", "documentos"),
                                         ("Desktop", "escritorio"),
                                         ("{374DE290-123F-4565-9164-39C4925E467B}", "descargas"),
                                         ("My Pictures", "imagenes")):
                    try:
                        crudo, _ = winreg.QueryValueEx(k, nombre)
                        p = Path(os.path.expandvars(crudo))
                        if p.exists():
                            sitios[etiqueta] = p
                    except OSError:
                        continue
        except Exception:
            pass

    # Respaldos, por si el registro no dijo nada.
    for etiqueta, p in (("documentos", casa / "Documents"),
                        ("escritorio", casa / "Desktop"),
                        ("descargas", casa / "Downloads"),
                        ("onedrive", casa / "OneDrive")):
        if etiqueta not in sitios and p.exists():
            sitios[etiqueta] = p

    if (casa / "OneDrive").exists():
        sitios.setdefault("onedrive", casa / "OneDrive")
    sitios["vault"] = VAULT
    sitios["colita"] = casa / "colita"
    sitios["casa"] = casa

    # Cualquier carpeta suelta del perfil que valga la pena: proyectos, código,
    # cursos... Diego trabaja en varias y no todas cuelgan de Documentos.
    for extra in ("Videos", "Music", "Pictures", "Proyectos", "PROYECTOS",
                  "repos", "dev", "Escritorio"):
        p = casa / extra
        if p.exists():
            sitios.setdefault(extra.lower(), p)

    # Otras unidades montadas (D:, E:, discos externos).
    if os.name == "nt":
        import string

        for letra in string.ascii_uppercase[3:]:      # de la D en adelante
            u = Path(f"{letra}:\\")
            try:
                if u.exists():
                    sitios[f"unidad {letra.lower()}"] = u
            except Exception:
                continue
    return sitios


CARPETAS = _carpetas_reales()

# Carpetas que nunca interesan al buscar: son entrañas del sistema y de los
# programas. Buscar en todo el perfil sin este filtro devuelve basura.
IGNORAR = frozenset({
    "AppData", "$Recycle.Bin", "Windows", "Program Files", "Program Files (x86)",
    ".git", "__pycache__", "node_modules", ".venv", "venv", "env",
    ".cache", ".conda", ".ipynb_checkpoints", ".vscode", ".gradle", ".nuget",
    "site-packages", "OneDriveTemp", "$WinREAgent", "System Volume Information",
})


def _ok(texto: str) -> dict[str, Any]:
    return {"content": [{"type": "text", "text": texto}]}


def _slug(texto: str) -> str:
    s = re.sub(r"[^\w\s-]", "", texto.lower()).strip()
    return re.sub(r"[\s_]+", "-", s)[:60] or "sin-titulo"


# ══════════════════════════════════════════════════════ sistema

@tool("volumen", "Sube, baja o silencia el volumen de Windows. accion: subir|bajar|silenciar",
      {"accion": str})
async def volumen(args: dict[str, Any]) -> dict[str, Any]:
    teclas = {"subir": 0xAF, "bajar": 0xAE, "silenciar": 0xAD}
    accion = str(args.get("accion", "")).lower()
    if accion not in teclas:
        return _ok("Acción no válida. Usa subir, bajar o silenciar.")
    import ctypes
    for _ in range(1 if accion == "silenciar" else 5):
        ctypes.windll.user32.keybd_event(teclas[accion], 0, 0, 0)
        ctypes.windll.user32.keybd_event(teclas[accion], 0, 2, 0)
    return _ok(f"Volumen: {accion}.")


@tool("abrir_app", "Abre una aplicación o archivo de Windows por su nombre o ruta.",
      {"que": str})
async def abrir_app(args: dict[str, Any]) -> dict[str, Any]:
    que = str(args.get("que", "")).strip()
    if not que:
        return _ok("Dime qué abrir.")
    try:
        os.startfile(que)          # nombre de app, ruta o URL
        return _ok(f"Abierto: {que}")
    except Exception:
        try:
            subprocess.Popen(["cmd", "/c", "start", "", que], shell=False, **SIN_VENTANA)
            return _ok(f"Abierto: {que}")
        except Exception as e:
            return _ok(f"No pude abrir «{que}»: {e}")


@tool(
    "poner_musica",
    "Pone una canción o vídeo de YouTube en el navegador normal de Diego. "
    "Es la forma correcta de poner música: NO uses Playwright para esto.",
    {"busqueda": str},
)
async def poner_musica(args: dict[str, Any]) -> dict[str, Any]:
    """Resuelve el primer resultado y lo abre en el navegador de siempre.

    Por que no con Playwright, que era lo que se hacia antes y se atascaba:

    - Playwright abre un perfil limpio, como una ventana de incognito. Ahi no
      esta la sesion de Diego ni su bloqueador, asi que salen todos los
      anuncios y hay que pelearse con ellos.
    - Cada intento de saltarlos pedia permiso para `browser_evaluate`, y la
      cosa acababa en un interrogatorio en vez de en musica.

    Abrirlo en el navegador de siempre resuelve las dos cosas de golpe: es su
    sesion, sus extensiones y sus suscripciones. Y no hace falta automatizar
    nada, que es la parte que se rompia.
    """
    import urllib.parse
    import urllib.request
    import webbrowser

    q = str(args.get("busqueda", "")).strip()
    if not q:
        return _ok("¿Qué quieres que ponga?")

    # sp=EgIQAQ== filtra a solo vídeos: evita caer en un canal o una playlist.
    busqueda = "https://www.youtube.com/results?" + urllib.parse.urlencode(
        {"search_query": q, "sp": "EgIQAQ=="}
    )
    destino, titulo = busqueda, None
    try:
        pet = urllib.request.Request(busqueda, headers={
            "User-Agent": ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                           "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0 Safari/537.36"),
            "Accept-Language": "es-PE,es;q=0.9",
        })
        with urllib.request.urlopen(pet, timeout=12) as r:
            html = r.read().decode("utf-8", "replace")
        ids = re.findall(r'"videoId":"([A-Za-z0-9_-]{11})"', html)
        if ids:
            destino = f"https://www.youtube.com/watch?v={ids[0]}"
            t = re.search(r'"title":\{"runs":\[\{"text":"(.*?)"\}', html)
            if t:
                # Con `unicode_escape` salía "TitÃ­ Me PreguntÃ³": ese códec
                # trata los bytes como latin-1. El texto viene con escapes
                # JSON (í), así que se decodifica con json.
                import json

                try:
                    titulo = json.loads(f'"{t.group(1)}"')
                except Exception:
                    titulo = t.group(1)
    except Exception:
        pass    # sin red o YouTube cambió el HTML: se abre la búsqueda y ya

    try:
        webbrowser.open(destino)
    except Exception as e:
        return _ok(f"No pude abrir el navegador: {e}")

    if titulo:
        return _ok(f"Sonando: {titulo}")
    if destino == busqueda:
        return _ok(f"No pude resolver el vídeo, así que te abrí la búsqueda de «{q}».")
    return _ok(f"Puesto lo primero que salió de «{q}».")


@tool(
    "buscar_archivo",
    "Busca archivos por nombre en las carpetas de Diego (Documentos, Escritorio, "
    "Descargas, OneDrive, vault y sus proyectos). Devuelve las rutas y cuándo se "
    "modificaron. Úsalo antes de decir que no encuentras algo.",
    {"nombre": str, "carpeta": str},
)
async def buscar_archivo(args: dict[str, Any]) -> dict[str, Any]:
    """Glob sabe buscar, pero hay que saber DÓNDE. Esto lleva el mapa dentro."""
    patron = str(args.get("nombre", "")).strip()
    if not patron:
        return _ok("¿Qué archivo busco?")
    # "ENAPRES*" exigiría que el nombre EMPIECE así, y los archivos de Diego se
    # llaman "TESIS MONTERO AGUA ENAPRES.docx". Nadie busca pensando en anclas:
    # se busca por "que contenga". Un patrón de extensión (*.pdf) se respeta.
    if not any(c in patron for c in "*?"):
        patron = f"*{patron}*"
    elif not patron.startswith("*"):
        patron = "*" + patron

    casa = Path.home()
    raices = [
        CARPETAS.get("documentos", casa / "Documents"),
        CARPETAS.get("escritorio", casa / "Desktop"),
        CARPETAS.get("descargas", casa / "Downloads"),
        casa / "Documents",                 # la de siempre, además de la redirigida
        VAULT, casa / "colita",
    ]
    if "onedrive" in CARPETAS:
        raices.append(CARPETAS["onedrive"])
    # Todo el perfil, no solo OneDrive. Es lo que Diego pidió el 2026-08-12:
    # que pueda mirar en cualquier carpeta suya, no en un par elegidas a mano.
    raices.append(casa)
    # Sin duplicados y sin carpetas contenidas en otra ya listada.
    unicas: list[Path] = []
    for r in raices:
        if r.exists() and not any(str(r).lower().startswith(str(u).lower() + os.sep)
                                  for u in unicas):
            if r not in unicas:
                unicas.append(r)
    raices = unicas

    pedida = str(args.get("carpeta", "")).strip()
    if pedida:
        p = CARPETAS.get(pedida.lower(), Path(pedida))
        if p.exists():
            raices = [p]

    vistos: list[tuple[float, Path]] = []
    for raiz in raices:
        if not raiz.exists():
            continue
        try:
            for f in raiz.rglob(patron):
                # Fuera el ruido: entrañas de programas, cachés y entornos.
                # Sin esto, buscar en todo el perfil devuelve miles de archivos
                # de AppData que a Diego no le sirven de nada.
                if any(x in f.parts for x in IGNORAR):
                    continue
                if f.is_file():
                    vistos.append((f.stat().st_mtime, f))
                    if len(vistos) > 400:
                        break
        except Exception:
            continue

    if not vistos:
        return _ok(f"No encontré nada que se parezca a «{patron}» en sus carpetas.")

    vistos.sort(reverse=True)              # lo más reciente primero: casi siempre es eso
    lineas = [
        f"{f}  ({dt.datetime.fromtimestamp(m):%Y-%m-%d %H:%M}, "
        f"{f.stat().st_size / 1024:.0f} KB)"
        for m, f in vistos[:15]
    ]
    extra = f"\n… y {len(vistos) - 15} más." if len(vistos) > 15 else ""
    return _ok(f"{len(vistos)} coincidencias, de más reciente a más antigua:\n"
               + "\n".join(lineas) + extra)


@tool(
    "listar_carpeta",
    "Muestra qué hay dentro de una carpeta de Diego: archivos, tamaños y fechas.",
    {"ruta": str},
)
async def listar_carpeta(args: dict[str, Any]) -> dict[str, Any]:
    ruta = str(args.get("ruta", "")).strip().strip('"')
    if not ruta:
        return _ok("¿Qué carpeta miro?")
    # Atajos por nombre, resueltos con las carpetas REALES de esta máquina.
    atajos = dict(CARPETAS)
    atajos.update({
        "documents": CARPETAS.get("documentos", Path.home() / "Documents"),
        "desktop": CARPETAS.get("escritorio", Path.home() / "Desktop"),
        "downloads": CARPETAS.get("descargas", Path.home() / "Downloads"),
        "obsidian": VAULT,
    })
    p = atajos.get(ruta.lower(), Path(ruta))
    if not p.exists():
        return _ok(f"No existe: {p}")
    if p.is_file():
        return _ok(f"{p} es un archivo de {p.stat().st_size / 1024:.0f} KB.")

    try:
        hijos = sorted(p.iterdir(), key=lambda x: (x.is_file(), x.name.lower()))
    except PermissionError:
        return _ok(f"Windows no me deja leer {p}.")

    lineas = []
    for h in hijos[:60]:
        try:
            m = dt.datetime.fromtimestamp(h.stat().st_mtime).strftime("%Y-%m-%d")
            if h.is_dir():
                lineas.append(f"[carpeta] {h.name}/   {m}")
            else:
                lineas.append(f"          {h.name}   {h.stat().st_size / 1024:.0f} KB   {m}")
        except Exception:
            continue
    extra = f"\n… y {len(hijos) - 60} más." if len(hijos) > 60 else ""
    return _ok(f"{p}  ({len(hijos)} elementos)\n" + "\n".join(lineas) + extra)


@tool(
    "abrir_carpeta",
    "Abre una carpeta en el Explorador de Windows, o selecciona un archivo dentro.",
    {"ruta": str},
)
async def abrir_carpeta(args: dict[str, Any]) -> dict[str, Any]:
    ruta = str(args.get("ruta", "")).strip().strip('"')
    p = Path(ruta)
    if not p.exists():
        return _ok(f"No existe: {p}")
    try:
        if p.is_file():
            subprocess.Popen(["explorer", "/select,", str(p)])
            return _ok(f"Abrí la carpeta con {p.name} seleccionado.")
        subprocess.Popen(["explorer", str(p)])
        return _ok(f"Abrí {p}.")
    except Exception as e:
        return _ok(f"No pude abrir el Explorador: {e}")


@tool(
    "leer_documento",
    "Lee un PDF, Word, Excel, CSV o texto y devuelve su contenido para poder "
    "resumirlo o analizarlo. Para PDF y Word extrae el texto.",
    {"ruta": str, "paginas": str},
)
async def leer_documento(args: dict[str, Any]) -> dict[str, Any]:
    """`Read` no abre PDF ni Word. Esto sí, que es la mitad de lo que le pasa."""
    ruta = str(args.get("ruta", "")).strip().strip('"')
    p = Path(ruta)
    if not p.exists():
        return _ok(f"No existe: {p}. Prueba a buscarlo con `buscar_archivo`.")

    ext = p.suffix.lower()
    TOPE = 60_000        # suficiente para resumir; más satura la conversación
    try:
        if ext == ".pdf":
            from pypdf import PdfReader

            lector = PdfReader(str(p))
            rango = str(args.get("paginas", "")).strip()
            hojas = range(len(lector.pages))
            if rango and "-" in rango:
                a, _, b = rango.partition("-")
                hojas = range(max(0, int(a) - 1), min(len(lector.pages), int(b)))
            texto = "\n".join(
                f"--- página {i + 1} ---\n{lector.pages[i].extract_text() or ''}"
                for i in hojas
            )
            cab = f"{p.name} · {len(lector.pages)} páginas\n\n"
        elif ext == ".docx":
            import docx

            d = docx.Document(str(p))
            texto = "\n".join(x.text for x in d.paragraphs)
            cab = f"{p.name} · {len(d.paragraphs)} párrafos\n\n"
        elif ext in {".xlsx", ".xls"}:
            import pandas as pd

            hojas = pd.read_excel(p, sheet_name=None)
            partes = []
            for nombre, df in hojas.items():
                partes.append(f"--- hoja «{nombre}» · {df.shape[0]}x{df.shape[1]} ---\n"
                              f"{df.head(20).to_string()}")
            texto = "\n\n".join(partes)
            cab = f"{p.name} · {len(hojas)} hojas\n\n"
        elif ext in {".csv", ".tsv"}:
            import pandas as pd

            df = pd.read_csv(p, sep=None, engine="python", nrows=200)
            texto = df.to_string()
            cab = f"{p.name} · columnas: {', '.join(map(str, df.columns))}\n\n"
        else:
            texto = p.read_text(encoding="utf-8", errors="replace")
            cab = f"{p.name}\n\n"
    except ImportError as e:
        return _ok(f"Me falta una librería para abrir {ext}: {e}")
    except Exception as e:
        return _ok(f"No pude leer {p.name}: {e}")

    if len(texto) > TOPE:
        texto = texto[:TOPE] + f"\n\n[…cortado, el documento sigue. Pídeme un rango de páginas.]"
    return _ok(cab + texto)


@tool(
    "portapapeles",
    "Lee lo que Diego tiene copiado, o copia algo para que él lo pegue. "
    "accion: leer|copiar. Es la forma más rápida de pasarle un texto largo.",
    {"accion": str, "texto": str},
)
async def portapapeles(args: dict[str, Any]) -> dict[str, Any]:
    accion = str(args.get("accion", "leer")).lower()
    try:
        import tkinter as tk

        raiz = tk.Tk()
        raiz.withdraw()
        try:
            if accion.startswith("cop"):
                texto = str(args.get("texto", ""))
                raiz.clipboard_clear()
                raiz.clipboard_append(texto)
                raiz.update()
                return _ok(f"Copiado ({len(texto)} caracteres). Pégalo con Ctrl+V.")
            contenido = raiz.clipboard_get()
        finally:
            raiz.destroy()
    except Exception as e:
        return _ok(f"No pude usar el portapapeles: {e}")

    if not contenido.strip():
        return _ok("El portapapeles está vacío.")
    if len(contenido) > 20_000:
        contenido = contenido[:20_000] + "\n[…cortado]"
    return _ok(f"Tienes copiado esto:\n\n{contenido}")


@tool(
    "recordar_luego",
    "Pon un recordatorio hablado. Colita lo dice en voz alta cuando toque. "
    "minutos: dentro de cuánto.",
    {"minutos": int, "recordatorio": str},
)
async def recordar_luego(args: dict[str, Any]) -> dict[str, Any]:
    """Un recordatorio que SUENA vale mucho más que uno que parpadea."""
    import threading

    try:
        minutos = max(1, min(720, int(args.get("minutos", 10))))
    except Exception:
        minutos = 10
    que = str(args.get("recordatorio", "")).strip()
    if not que:
        return _ok("¿Qué te recuerdo?")

    def avisar() -> None:
        try:
            import sys

            sys.path.insert(0, str(Path(__file__).parent))
            import voz

            voz.hablar(f"Diego, me pediste que te recordara: {que}", bloquear=False)
        except Exception:
            pass

    t = threading.Timer(minutos * 60, avisar)
    t.daemon = True
    t.start()
    cuando = (dt.datetime.now() + dt.timedelta(minutes=minutos)).strftime("%H:%M")
    return _ok(f"Listo. Te lo digo a las {cuando}: «{que}».")


@tool(
    "resumen_del_dia",
    "Qué se movió hoy en la máquina de Diego: notas nuevas o editadas del vault, "
    "archivos que tocó y estado del equipo. Para cerrar o abrir la jornada.",
    {"dias": int},
)
async def resumen_del_dia(args: dict[str, Any]) -> dict[str, Any]:
    try:
        dias = max(1, min(30, int(args.get("dias", 1))))
    except Exception:
        dias = 1
    corte = dt.datetime.now() - dt.timedelta(days=dias)
    marca = corte.timestamp()

    notas = []
    for p in VAULT.rglob("*.md"):
        try:
            if p.stat().st_mtime > marca:
                notas.append((p.stat().st_mtime, p))
        except Exception:
            continue
    notas.sort(reverse=True)

    archivos = []
    for raiz in (CARPETAS.get("descargas"), CARPETAS.get("documentos"),
                 CARPETAS.get("escritorio")):
        if raiz is None or not raiz.exists():
            continue
        try:
            for f in raiz.rglob("*"):
                if any(x in f.parts for x in IGNORAR):
                    continue
                if f.is_file() and f.stat().st_mtime > marca:
                    archivos.append((f.stat().st_mtime, f))
                    if len(archivos) > 200:
                        break
        except Exception:
            continue
    archivos.sort(reverse=True)

    partes = [f"En {'el último día' if dias == 1 else f'los últimos {dias} días'}:"]
    if notas:
        partes.append(f"\n{len(notas)} notas tocadas en el vault:")
        partes += [f"  · {p.stem}" for _, p in notas[:12]]
        if len(notas) > 12:
            partes.append(f"  … y {len(notas) - 12} más.")
    else:
        partes.append("\nNinguna nota nueva en el vault.")

    if archivos:
        partes.append(f"\n{len(archivos)} archivos tocados:")
        partes += [f"  · {f.name}   ({f.parent})" for _, f in archivos[:10]]
    return _ok("\n".join(partes))


@tool("estado_maquina", "Uso de CPU, memoria y disco de la laptop.", {})
async def estado_maquina(args: dict[str, Any]) -> dict[str, Any]:
    import psutil
    cpu = psutil.cpu_percent(interval=0.4)
    m = psutil.virtual_memory()
    d = psutil.disk_usage("C:\\")
    return _ok(
        f"CPU {cpu:.0f}% | RAM {m.percent:.0f}% ({m.used/1e9:.1f} de {m.total/1e9:.1f} GB) "
        f"| Disco C: {d.percent:.0f}% ({d.free/1e9:.0f} GB libres)"
    )


# ══════════════════════════════════════════════════════ cerebro (vault)

@tool(
    "guardar_en_vault",
    "Guarda conocimiento nuevo en el vault de Obsidian de Diego, para que no se pierda. "
    "carpeta: metodos|conceptos|aprendizajes|papers|proyectos|snippets",
    {"titulo": str, "carpeta": str, "contenido": str, "tags": str},
)
async def guardar_en_vault(args: dict[str, Any]) -> dict[str, Any]:
    carpetas = {"metodos", "conceptos", "aprendizajes", "papers", "proyectos", "snippets"}
    carpeta = str(args.get("carpeta", "conceptos")).lower()
    if carpeta not in carpetas:
        carpeta = "conceptos"

    titulo = str(args.get("titulo", "")).strip()
    if not titulo:
        return _ok("Necesito un título para la nota.")

    tipo = {"metodos": "metodo", "conceptos": "concepto", "aprendizajes": "aprendizaje",
            "papers": "paper", "proyectos": "proyecto", "snippets": "snippet"}[carpeta]
    tags = str(args.get("tags", "")).strip()
    lista = ", ".join(t.strip() for t in tags.split(",") if t.strip()) or tipo
    hoy = dt.date.today().isoformat()

    ruta = VAULT / carpeta / f"{_slug(titulo)}.md"
    if ruta.exists():
        return _ok(f"Ya existe esa nota: {ruta.name}. Léela y actualízala en vez de duplicar.")

    ruta.write_text(
        f"---\ntipo: {tipo}\ncreado: {hoy}\ntags: [{lista}]\n---\n\n"
        f"# {titulo}\n\n{args.get('contenido', '').strip()}\n",
        encoding="utf-8",
    )
    return _ok(f"Guardado en {carpeta}/{ruta.name}. Recuerda enlazarlo desde el MOC que toque.")


@tool(
    "enlazar_en_moc",
    "Añade un enlace a una nota dentro de un MOC, bajo la sección que se indique. "
    "Sin esto, cada nota nueva queda huérfana y no se encuentra navegando. "
    "moc: el nombre del archivo sin .md, p. ej. moc-ciclo-10-unmsm",
    {"moc": str, "nota": str, "descripcion": str, "seccion": str},
)
async def enlazar_en_moc(args: dict[str, Any]) -> dict[str, Any]:
    """Cierra el bucle de `guardar_en_vault`.

    Antes, Colita creaba la nota y decia "recuerda enlazarla desde el MOC" —
    dejandole a Diego el trabajo aburrido, que es justo el que se olvida. Una
    nota sin enlaces entrantes es una nota que no existe.
    """
    moc = _slug(str(args.get("moc", "")).replace(".md", ""))
    nota = _slug(str(args.get("nota", "")).replace(".md", ""))
    if not moc or not nota:
        return _ok("Necesito el MOC y la nota.")

    destino = next((p for p in VAULT.rglob(f"{moc}.md")), None)
    if destino is None:
        return _ok(f"No encuentro el MOC '{moc}'. Mira los que hay antes de inventarlo.")
    if not any(VAULT.rglob(f"{nota}.md")):
        return _ok(f"La nota '{nota}' no existe todavía. Créala primero.")

    texto = destino.read_text(encoding="utf-8")
    if f"[[{nota}]]" in texto or f"[[{nota}|" in texto:
        return _ok(f"{nota} ya estaba enlazada en {destino.name}. No toqué nada.")

    desc = str(args.get("descripcion", "")).strip()
    linea = f"- [[{nota}]]" + (f" — {desc}" if desc else "")

    seccion = str(args.get("seccion", "")).strip()
    lineas = texto.splitlines()
    puesto = None
    if seccion:
        # Debajo del encabezado pedido, al final de su bloque.
        for i, l in enumerate(lineas):
            if l.startswith("#") and seccion.lower() in l.lower():
                j = i + 1
                while j < len(lineas) and not lineas[j].startswith("#"):
                    j += 1
                while j > i + 1 and not lineas[j - 1].strip():
                    j -= 1               # sin dejar huecos raros
                puesto = j
                break
    if puesto is None:
        # Sin sección: antes del "Ver también" final, que siempre cierra la nota.
        for i in range(len(lineas) - 1, -1, -1):
            if lineas[i].startswith("Ver también"):
                puesto = i
                while puesto > 0 and not lineas[puesto - 1].strip():
                    puesto -= 1
                break
    if puesto is None:
        puesto = len(lineas)

    lineas.insert(puesto, linea)
    destino.write_text("\n".join(lineas) + "\n", encoding="utf-8")
    return _ok(f"Enlazada en {destino.name}: {linea}")


@tool(
    "salud_del_vault",
    "Revisa el vault entero: cuántas notas hay, enlaces rotos y notas huérfanas "
    "(sin ningún enlace entrante). Úsalo tras añadir varias notas.",
    {},
)
async def salud_del_vault(args: dict[str, Any]) -> dict[str, Any]:
    """Un enlace roto o una nota huérfana no dan error: solo se pierden."""
    notas = list(VAULT.rglob("*.md"))
    existentes = {p.stem for p in notas}
    entrantes: dict[str, int] = {s: 0 for s in existentes}
    rotos: list[str] = []

    for p in notas:
        # Las plantillas llevan marcadores a propósito.
        if p.parent.name == "plantillas":
            continue
        for e in re.findall(r"\[\[([^\]|#]+)", p.read_text(encoding="utf-8")):
            e = e.strip()
            if e in existentes:
                entrantes[e] += 1
            else:
                rotos.append(f"[[{e}]] en {p.name}")

    huerfanas = sorted(s for s, n in entrantes.items() if n == 0)
    partes = [f"{len(notas)} notas."]
    if rotos:
        partes.append(f"{len(rotos)} enlaces rotos (los primeros): " + "; ".join(rotos[:8]))
        partes.append("Ojo: un enlace a una nota que aún no existe es válido si es un "
                      "hueco declarado a propósito.")
    else:
        partes.append("Sin enlaces rotos.")
    if huerfanas:
        partes.append(f"{len(huerfanas)} notas sin enlaces entrantes: "
                      + ", ".join(huerfanas[:10]))
    else:
        partes.append("Ninguna nota huérfana.")
    return _ok(" ".join(partes))


@tool(
    "nota_de_clase",
    "Crea la nota de una clase del ciclo en curso, con su plantilla, y la enlaza sola "
    "desde el MOC del ciclo. curso: el nombre tal cual, p. ej. Análisis Bayesiano",
    {"curso": str, "tema": str, "contenido": str},
)
async def nota_de_clase(args: dict[str, Any]) -> dict[str, Any]:
    curso = str(args.get("curso", "")).strip()
    tema = str(args.get("tema", "")).strip()
    if not curso or not tema:
        return _ok("Necesito el curso y el tema de la clase.")

    carpeta = VAULT / "cursos"
    carpeta.mkdir(exist_ok=True)
    hoy = dt.date.today().isoformat()
    nombre = _slug(f"{curso}-{hoy}-{tema}")
    ruta = carpeta / f"{nombre}.md"
    if ruta.exists():
        return _ok(f"Ya existe {ruta.name}. Ábrela y añade ahí en vez de duplicar.")

    ruta.write_text(
        f"---\ntipo: clase\ncurso: {curso}\nestado: en-curso\ncreado: {hoy}\n"
        f"tags: [clase, ciclo-10, unmsm]\n---\n\n"
        f"# {curso} — {tema}\n\n"
        f"Clase del `{hoy}`.\n\n"
        f"## Lo que se vio\n\n{args.get('contenido', '').strip()}\n\n"
        f"## Lo que no me quedó claro\n\n- \n\n"
        f"## Qué tengo que practicar\n\n- \n\n"
        f"Ver también: [[moc-ciclo-10-unmsm]] · [[malla-curricular-estadistica-unmsm]]\n",
        encoding="utf-8",
    )
    return _ok(f"Creada cursos/{ruta.name}. Dime si quieres que la enlace desde el MOC "
               f"del ciclo con `enlazar_en_moc`.")


@tool("reindexar_vault", "Reconstruye el índice de búsqueda del vault tras añadir notas.", {})
async def reindexar_vault(args: dict[str, Any]) -> dict[str, Any]:
    uvx = Path(os.environ["LOCALAPPDATA"]) / "Microsoft/WinGet/Links/uvx.exe"
    entorno = {
        **os.environ,
        "OBSIDIAN_RAG_VAULT": str(VAULT),
        "OBSIDIAN_RAG_PROVIDER": "ollama",
        "OBSIDIAN_RAG_MODEL": "nomic-embed-text",
        "OBSIDIAN_RAG_OLLAMA_URL": "http://127.0.0.1:11434",
    }
    try:
        r = subprocess.run(
            [str(uvx), "--with", "mcp<2", "obsidian-notes-rag", "index"],
            env=entorno, capture_output=True, text=True, timeout=900, **SIN_VENTANA,
        )
        linea = next((l for l in r.stdout.splitlines() if "Indexed" in l), "reindexado")
        return _ok(linea.strip())
    except Exception as e:
        return _ok(f"No pude reindexar: {e}")


# ══════════════════════════════════════════════════════ entregables

@tool(
    "crear_excel",
    "Crea un .xlsx a partir de datos. filas: lista de listas en JSON, la primera es la cabecera.",
    {"nombre": str, "filas_json": str, "hoja": str},
)
async def crear_excel(args: dict[str, Any]) -> dict[str, Any]:
    import json
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill

    try:
        filas = json.loads(str(args.get("filas_json", "[]")))
    except Exception as e:
        return _ok(f"El JSON de filas no es válido: {e}")
    if not filas:
        return _ok("No me diste datos.")

    wb = Workbook()
    ws = wb.active
    ws.title = str(args.get("hoja", "Datos"))[:31]
    for f in filas:
        ws.append(list(f))

    cab = Font(bold=True, color="FFFFFF")
    relleno = PatternFill("solid", fgColor="1F3B5C")
    for c in ws[1]:
        c.font, c.fill = cab, relleno
    for col in ws.columns:
        ancho = max((len(str(c.value)) for c in col if c.value), default=8)
        ws.column_dimensions[col[0].column_letter].width = min(ancho + 3, 55)
    ws.freeze_panes = "A2"

    ruta = SALIDAS / f"{_slug(str(args.get('nombre', 'datos')))}.xlsx"
    wb.save(ruta)
    return _ok(f"Excel creado: {ruta}")


@tool(
    "crear_html",
    "Crea un informe HTML autocontenido, oscuro y legible. El cuerpo va en HTML.",
    {"titulo": str, "cuerpo_html": str},
)
async def crear_html(args: dict[str, Any]) -> dict[str, Any]:
    titulo = str(args.get("titulo", "Informe")).strip()
    ruta = SALIDAS / f"{_slug(titulo)}.html"
    ruta.write_text(
        "<!doctype html><html lang=es><meta charset=utf-8>"
        f"<title>{titulo}</title><style>"
        ":root{color-scheme:dark}"
        "body{background:#0a0e16;color:#dce6f5;font:16px/1.7 'Segoe UI',system-ui,sans-serif;"
        "max-width:860px;margin:0 auto;padding:48px 24px}"
        "h1{font-size:1.9rem;letter-spacing:-.02em;background:linear-gradient(92deg,#00e5ff,#8b5cff);"
        "-webkit-background-clip:text;background-clip:text;color:transparent}"
        "h2{margin-top:2.2em;color:#00e5ff;font-size:1.25rem}"
        "table{border-collapse:collapse;width:100%;margin:1.4em 0}"
        "th,td{border:1px solid #1e2a40;padding:9px 13px;text-align:left}"
        "th{background:#101a2b;color:#8fb4d9}tr:nth-child(even) td{background:#0d1420}"
        "code{background:#101a2b;padding:2px 6px;border-radius:4px;color:#00ff9d}"
        "pre{background:#0d1420;padding:16px;border-radius:10px;overflow-x:auto;"
        "border:1px solid #1e2a40}"
        "a{color:#00e5ff}.pie{margin-top:3em;color:#5b6b84;font-size:.85rem;"
        "border-top:1px solid #1e2a40;padding-top:1em}"
        f"</style><h1>{titulo}</h1>{args.get('cuerpo_html', '')}"
        f"<p class=pie>Generado por Colita AI · {dt.date.today().isoformat()}</p></html>",
        encoding="utf-8",
    )
    return _ok(f"Informe creado: {ruta}")


@tool(
    "analizar_datos",
    "Ejecuta código Python de análisis (pandas, numpy, sklearn, matplotlib) y devuelve la salida. "
    "Usa print() para lo que quieras ver. Las figuras se guardan en salidas/.",
    {"codigo": str},
)
async def analizar_datos(args: dict[str, Any]) -> dict[str, Any]:
    import sys
    import textwrap

    codigo = textwrap.dedent(str(args.get("codigo", "")))
    envoltorio = (
        "import matplotlib\nmatplotlib.use('Agg')\n"
        "import matplotlib.pyplot as plt, pandas as pd, numpy as np, os\n"
        f"os.chdir(r'{SALIDAS}')\n" + codigo + "\n"
        "for i in plt.get_fignums():\n"
        "    plt.figure(i).savefig(f'figura_{i}.png', dpi=140, bbox_inches='tight')\n"
    )
    try:
        r = subprocess.run(
            [sys.executable, "-c", envoltorio],
            capture_output=True, text=True, timeout=300, cwd=str(SALIDAS),
            **SIN_VENTANA,
        )
    except subprocess.TimeoutExpired:
        return _ok("El análisis tardó más de 5 minutos y lo corté.")

    salida = (r.stdout or "").strip()
    error = (r.stderr or "").strip()
    if error and r.returncode != 0:
        return _ok(f"Falló:\n{error[-1500:]}")
    return _ok(salida[-4000:] or "Se ejecutó sin imprimir nada.")


# ══════════════════════════════════════════════════════

servidor = create_sdk_mcp_server(
    name="colita",
    version="1.0.0",
    tools=[
        volumen, abrir_app, poner_musica, estado_maquina,
        portapapeles, recordar_luego, resumen_del_dia,
        buscar_archivo, listar_carpeta, abrir_carpeta, leer_documento,
        guardar_en_vault, enlazar_en_moc, salud_del_vault, nota_de_clase,
        reindexar_vault,
        crear_excel, crear_html, analizar_datos,
    ],
)

NOMBRES = [
    "mcp__colita__volumen", "mcp__colita__abrir_app", "mcp__colita__poner_musica",
    "mcp__colita__estado_maquina",
    "mcp__colita__portapapeles", "mcp__colita__recordar_luego",
    "mcp__colita__resumen_del_dia",
    "mcp__colita__buscar_archivo", "mcp__colita__listar_carpeta",
    "mcp__colita__abrir_carpeta", "mcp__colita__leer_documento",
    "mcp__colita__guardar_en_vault", "mcp__colita__enlazar_en_moc",
    "mcp__colita__salud_del_vault", "mcp__colita__nota_de_clase",
    "mcp__colita__reindexar_vault",
    "mcp__colita__crear_excel", "mcp__colita__crear_html", "mcp__colita__analizar_datos",
]
